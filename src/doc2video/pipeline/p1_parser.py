"""P1 文档解析:PDF/DOCX/MD/TXT → parsed.json(含阅读顺序/坐标/OCR 逐页判定)。

设计要点(方案 4.2):
- PDF: TOC 优先分节;block/span 坐标 + 字号/加粗判定层级;多栏按坐标重排;页眉页脚去重;
  表格区域识别;扫描页逐页判定 → rapidocr;加密拒绝;parser_version 变化须级联失效下游。
- DOCX: 样式名 + outline level 映射;列表/表格/内嵌图片;不支持元素记入 meta.notes。
- block_id 稳定构造:全局序号 b{n}(解析顺序确定 ⇒ 同版本同输入必稳定)。
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from ..contracts import Block, ParsedDocument, ParsedMeta, Section
from .p0_ingest import RejectError

PARSER_VERSION = "0.1.0"

# ── 公共工具 ────────────────────────────────────────────────────


class _Builder:
    """块收集器:统一编号、分节、产出 ParsedDocument。"""

    def __init__(self, title: str, meta: dict[str, Any]):
        self.title = title
        self.meta = meta
        self.sections: list[Section] = []
        self._counter = 0

    def new_block(self, btype: str, text: str, page: int | None = None, bbox: list[float] | None = None,
                  reading_order: int | None = None, ocr_confidence: float | None = None,
                  section: Section | None = None) -> Block:
        self._counter += 1
        block = Block(
            block_id=f"b{self._counter}", type=btype, text=text, page=page, bbox=bbox,
            reading_order=reading_order, ocr_confidence=ocr_confidence,
        )
        target = section or self.sections[-1]
        target.blocks.append(block)
        return block

    def add_section(self, level: int, heading: str) -> Section:
        sec = Section(id=f"s{len(self.sections) + 1}", level=level, heading=heading)
        self.sections.append(sec)
        return sec

    def build(self, parser_notes: list[str] | None = None) -> ParsedDocument:
        if not self.sections:
            self.add_section(1, "正文")
        meta = ParsedMeta(parser_version=PARSER_VERSION, notes=parser_notes or [], **self.meta)
        return ParsedDocument(meta=meta, title=self.title, sections=self.sections)


# ── PDF ─────────────────────────────────────────────────────────


def _span_bold(flags: int) -> bool:
    return bool(flags & 2**4)


def _detect_columns(spans: list[dict]) -> list[tuple[float, float]]:
    """按 x 中心聚类出列区间(绝对坐标;单列 → 全宽,双列 → 左/右两带)。"""
    xs = sorted((s["bbox"][0] + s["bbox"][2]) / 2 for s in spans)
    if len(xs) < 8:
        return [(0.0, float("inf"))]
    lo, hi = xs[0], xs[-1]
    width = hi - lo
    # 找中点附近 ±0.15width 的"无字带"
    mid = (lo + hi) / 2
    gap = [x for x in xs if mid - 0.15 * width < x < mid + 0.15 * width]
    if len(gap) <= 0.02 * len(xs):  # 中间几乎没有字 → 双栏
        return [(0.0, mid), (mid, float("inf"))]
    return [(0.0, float("inf"))]


def _in_any_table(bbox: list[float], tables: list[Any]) -> bool:
    x0, y0, x1, y1 = bbox
    for t in tables:
        try:
            tx0, ty0, tx1, ty1 = t.bbox
        except AttributeError:
            continue
        if x0 >= tx0 - 2 and y0 >= ty0 - 2 and x1 <= tx1 + 2 and y1 <= ty1 + 2:
            return True
    return False


def _extract_table_blocks(page: Any, builder: _Builder, pno: int, section: Section) -> None:
    try:
        tables = page.find_tables()
    except Exception:  # noqa: BLE001 —— 无表格或识别失败均静默
        return
    for t in tables.tables:
        rows = []
        for row in t.extract():
            rows.append([(c or "").strip().replace("\n", " ") for c in row])
        if rows:
            text = " | ".join(" ; ".join(r) for r in rows)
            builder.new_block("table", text, page=pno,
                              bbox=list(t.bbox) if hasattr(t, "bbox") else None,
                              section=section)


def _dedup_header_footer(page_texts: list[set[str]], text: str) -> bool:
    """跨页重复的页眉页脚文本(≥3 页出现)→ 剔除。"""
    if len(page_texts) < 3:
        return False
    return sum(1 for s in page_texts if text in s) >= 3


def _parse_pdf(src: Path, job_dir: Path, cfg: dict[str, Any]) -> ParsedDocument:
    import pymupdf

    pdf = pymupdf.open(str(src))
    try:
        if pdf.needs_pass:
            raise RejectError("REJECT_ENCRYPTED: 加密 PDF 不支持")
        max_pages = cfg["limits"]["max_pdf_pages"]
        if pdf.page_count > max_pages:
            raise RejectError(f"REJECT_TOO_LARGE: 页数超限({pdf.page_count} > {max_pages})")

        title = Path(src).stem
        builder = _Builder(title, {
            "source": str(src), "type": "pdf", "pages": pdf.page_count,
            "chars": sum(len(p.get_text("text")) for p in pdf),
        })

        toc = pdf.get_toc()
        toc_sections: list[tuple[int, str, int]] = [(lvl, t, pg) for lvl, t, pg in toc]
        section_by_page: dict[int, Section] = {}

        page_texts: list[set[str]] = []
        for page in pdf:
            page_texts.append({s.strip() for s in page.get_text("text").splitlines() if s.strip()})

        min_chars = cfg["parser"]["ocr_min_chars_per_page"]
        cur_section: Section | None = None
        for pno in range(1, pdf.page_count + 1):
            page = pdf[pno - 1]
            # TOC/默认分节(先于扫描判定,保证 OCR 块有归属节)
            for lvl, heading, tpage in toc_sections:
                if tpage == pno:
                    cur_section = builder.add_section(min(lvl, 6), heading)
                    section_by_page[pno] = cur_section
            if not builder.sections:
                cur_section = builder.add_section(1, "正文")
                section_by_page[pno] = cur_section
            active = section_by_page.get(pno) or cur_section or builder.sections[-1]

            # 扫描页判定:逐页,文本层过薄 → OCR
            plain = page.get_text("text")
            if len(plain.strip()) < min_chars:
                _ocr_page(page, builder, pno, active)
                continue

            d = page.get_text("dict")
            spans = [
                {"text": s["text"], "bbox": s["bbox"], "size": s["size"], "flags": s["flags"]}
                for b in d.get("blocks", []) if b.get("type") == 0
                for l in b.get("lines", []) for s in l.get("spans", [])
                if s.get("text", "").strip()
            ]
            if not spans:
                continue
            tables = []
            try:
                tables = page.find_tables().tables
            except Exception:  # noqa: BLE001, S110 —— 无表格/不支持时静默
                pass
            _extract_table_blocks(page, builder, pno, active)
            spans = [s for s in spans if not _in_any_table(s["bbox"], tables)]

            columns = _detect_columns(spans)
            fonts = [s["size"] for s in spans]
            median = sorted(fonts)[len(fonts) // 2] if fonts else 10.0
            order = 0
            for col_lo, col_hi in columns:
                col_spans = [
                    s for s in spans if col_lo <= (s["bbox"][0] + s["bbox"][2]) / 2 < col_hi
                ]
                col_spans.sort(key=lambda s: (round(s["bbox"][1] / 3), s["bbox"][0]))
                # 行分组 → 段落块
                lines: list[list[dict]] = []
                for s in col_spans:
                    y = round(s["bbox"][1] / 3)
                    if lines and abs(y - round(lines[-1][0]["bbox"][1] / 3)) <= 1:
                        lines[-1].append(s)
                    else:
                        lines.append([s])
                para: list[dict] = []
                for line in lines:
                    line_text = "".join(x["text"] for x in line)
                    if _dedup_header_footer(page_texts, line_text):
                        continue
                    is_heading = any(_span_bold(x["flags"]) for x in line) or any(
                        x["size"] > median * 1.18 for x in line
                    )
                    if is_heading:
                        if para:
                            order += 1
                            builder.new_block(
                                "paragraph", "".join(x["text"] for x in para),
                                page=pno, reading_order=order,
                                bbox=[para[0]["bbox"][0], para[0]["bbox"][1],
                                      para[-1]["bbox"][2], para[-1]["bbox"][3]],
                                section=active,
                            )
                            para = []
                        order += 1
                        builder.new_block(
                            "paragraph", line_text, page=pno, reading_order=order,
                            bbox=list(line[0]["bbox"]), section=active,
                        )
                    else:
                        para.extend(line)
                if para:
                    order += 1
                    builder.new_block(
                        "paragraph", "".join(x["text"] for x in para), page=pno,
                        reading_order=order,
                        bbox=[para[0]["bbox"][0], para[0]["bbox"][1],
                              para[-1]["bbox"][2], para[-1]["bbox"][3]],
                        section=active,
                    )
        return builder.build()
    finally:
        pdf.close()


def _ocr_page(page: Any, builder: _Builder, pno: int, section: Section | None = None) -> None:
    """扫描页:渲染 → rapidocr → blocks(带 ocr_confidence)。"""
    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError as exc:
        raise RejectError(f"REJECT_MALFORMED: OCR 引擎不可用(rapidocr-onnxruntime): {exc}")
    pix = page.get_pixmap(dpi=150)
    import numpy as np

    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    if pix.n == 4:
        img = img[:, :, :3]
    ocr = RapidOCR()
    result, _ = ocr(img)
    if not result:
        builder.new_block("paragraph", "", page=pno, ocr_confidence=0.0, section=section)
        return
    lines: list[tuple[float, str, float]] = []
    for box, text, conf in result:
        ys = [p[1] for p in box]
        lines.append((min(ys), str(text), float(conf)))
    lines.sort()
    for y, text, conf in lines:
        builder.new_block("paragraph", text, page=pno, bbox=None, ocr_confidence=conf,
                          section=section)


# ── DOCX ───────────────────────────────────────────────────────


def _para_is_list(para: Any) -> bool:
    """列表判定:直接 numPr 或列表类样式(List Bullet/Number,含样式继承)。"""
    pPr = para._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return True
    style = (para.style.name or "") if para.style is not None else ""
    s = style.lower()
    return s.startswith("list") or "bullet" in s or "number" in s


def _heading_level(style_name: str, outline: int | None) -> int | None:
    if outline is not None and 0 <= outline <= 8:
        return min(outline + 1, 6)
    m = re.search(r"(Heading|标题)\s*(\d)", style_name, re.IGNORECASE)
    if m:
        return min(int(m.group(2)), 6)
    if style_name.lower().startswith("title"):
        return 1
    return None


def _parse_docx(src: Path, job_dir: Path, cfg: dict[str, Any]) -> ParsedDocument:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(src))
    builder = _Builder(Path(src).stem, {
        "source": str(src), "type": "docx", "pages": 1,  # DOCX 页数依赖排版,解析期不可知 → 记 1
        "chars": sum(len(p.text) for p in doc.paragraphs),
    })
    notes: list[str] = []
    cur: Section | None = None

    def ensure_section(level: int, heading: str) -> Section:
        nonlocal cur
        if cur is None or heading:
            cur = builder.add_section(level, heading or "正文")
        return cur

    def flush_items(items: list[str]) -> None:
        if items:
            builder.new_block("list", "\n".join(items))

    body = doc.element.body
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    pending_items: list[str] = []
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            para = Paragraph(child, doc)
            text = para.text.strip()
            if not text:
                continue
            lvl = None
            try:
                outline = para.paragraph_format.outline_level
            except Exception:  # noqa: BLE001
                outline = None
            style = para.style.name if para.style is not None else ""
            lvl = _heading_level(style, outline)
            if lvl is not None:
                flush_items(pending_items)
                cur = ensure_section(lvl, text)
            elif _para_is_list(para):
                pending_items.append(text)
            else:
                flush_items(pending_items)
                ensure_section(1, "")
                builder.new_block("paragraph", text)
        elif child.tag.endswith("}tbl"):
            flush_items(pending_items)
            ensure_section(1, "")
            table = Table(child, doc)
            rows = []
            for row in table.rows:
                rows.append([c.text.strip() for c in row.cells])
            if rows:
                builder.new_block("table", " | ".join(" ; ".join(r) for r in rows))
        elif child.tag.endswith("}sdt"):
            notes.append("含结构化文档标签(SmartArt/控件),内容未解析")

    flush_items(pending_items)

    # 内嵌图片 → extracted_assets
    assets = job_dir / "extracted_assets"
    assets.mkdir(parents=True, exist_ok=True)
    for rel in doc.part.rels.values():
        if "image" in (rel.reltype or ""):
            try:
                data = rel.target_part.blob
                name = f"docx_{rel.rId}.{rel.target_part.partname.rsplit('.', 1)[-1]}"
                (assets / name).write_bytes(data)
                builder.new_block("image", f"extracted_assets/{name}")
            except Exception:  # noqa: BLE001
                notes.append(f"图片提取失败: {rel.rId}")
    return builder.build(parser_notes=notes)


# ── MD / TXT ──────────────────────────────────────────────────


def _parse_markdown(src: Path, job_dir: Path, cfg: dict[str, Any]) -> ParsedDocument:
    import markdown_it

    md = markdown_it.MarkdownIt("commonmark")
    md.enable("table")
    text = src.read_text(encoding="utf-8")
    builder = _Builder(Path(src).stem, {"source": str(src), "type": "md", "pages": 1, "chars": len(text)})
    builder.add_section(1, "正文")
    tokens = md.parse(text)
    i = 0
    while i < len(tokens):
        t = tokens[i]
        if t.type == "heading_open":
            level = int(t.tag[1])
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            heading = inline.content if inline and inline.type == "inline" else ""
            builder.add_section(level, heading)
            i += 3
            continue
        if t.type == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            builder.new_block("paragraph", inline.content if inline else "")
            i += 3
            continue
        if t.type == "bullet_list_open" or t.type == "ordered_list_open":
            items: list[str] = []
            j = i + 1
            while j < len(tokens) and tokens[j].type not in ("bullet_list_close", "ordered_list_close"):
                if tokens[j].type == "list_item_open":
                    k = j + 1
                    while k < len(tokens) and tokens[k].type not in ("inline", "list_item_close"):
                        k += 1
                    items.append(tokens[k].content if k < len(tokens) and tokens[k].type == "inline" else "")
                j += 1
            builder.new_block("list", "\n".join(items))
            i = j + 1
            continue
        if t.type == "table_open":
            rows = []
            j = i + 1
            while j < len(tokens) and tokens[j].type != "table_close":
                if tokens[j].type == "tr_open":
                    cells = []
                    k = j + 1
                    while k < len(tokens) and tokens[k].type != "tr_close":
                        if tokens[k].type == "td_open" or tokens[k].type == "th_open":
                            cells.append(tokens[k + 1].content if k + 1 < len(tokens) else "")
                        k += 1
                    rows.append(cells)
                j += 1
            builder.new_block("table", " | ".join(" ; ".join(r) for r in rows))
            i = j + 1
            continue
        i += 1
    return builder.build()


def _parse_txt(src: Path, job_dir: Path, cfg: dict[str, Any]) -> ParsedDocument:
    text = src.read_text(encoding="utf-8", errors="replace")
    builder = _Builder(Path(src).stem, {"source": str(src), "type": "txt", "pages": 1, "chars": len(text)})
    builder.add_section(1, "正文")
    for para in re.split(r"\n\s*\n", text):
        para = para.strip()
        if para:
            builder.new_block("paragraph", para)
    return builder.build()


# ── 入口 ───────────────────────────────────────────────────────


def parse_document(src: Path, job_dir: Path, doc_type: str, cfg: dict[str, Any]) -> ParsedDocument:
    if doc_type == "pdf":
        return _parse_pdf(src, job_dir, cfg)
    if doc_type == "docx":
        return _parse_docx(src, job_dir, cfg)
    if doc_type == "md":
        return _parse_markdown(src, job_dir, cfg)
    if doc_type == "txt":
        return _parse_txt(src, job_dir, cfg)
    raise RejectError(f"REJECT_UNSUPPORTED: 未知类型 {doc_type}")


def stage_p1(job_dir: Path, cfg: dict[str, Any], opts: Any, stage: str | None = None) -> Any:
    """P1 阶段入口:manifest → parsed.json + extracted_assets/。"""
    from ..contracts import Manifest
    from ..state import atomic_write_text
    from .stages import StageResult

    manifest = Manifest.model_validate_json((job_dir / "manifest.json").read_text(encoding="utf-8"))
    src = job_dir / "input" / Path(manifest.source).name
    doc = parse_document(src, job_dir, manifest.doc_type, cfg)
    atomic_write_text(job_dir / "parsed.json", doc.model_dump_json(indent=2) + "\n")
    artifacts = [("parsed.json", "application/json")]
    assets_dir = job_dir / "extracted_assets"
    if assets_dir.exists():
        for f in sorted(assets_dir.iterdir()):
            artifacts.append((f"extracted_assets/{f.name}", "application/octet-stream"))
    return StageResult(artifacts=artifacts)
