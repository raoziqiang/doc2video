"""M1 测试夹具工厂:各格式样例文档生成(纯内存/临时文件,不依赖外部样例)。"""

from __future__ import annotations

import zipfile
from pathlib import Path

BODY_ZH = "咖啡因是世界上最常用的精神活性物质,它通过阻断腺苷受体让人保持清醒。"


def make_md(tmp: Path) -> Path:
    p = tmp / "demo.md"
    p.write_text(
        "# 咖啡因与睡眠\n\n"
        "## 第一节 机制\n\n"
        "咖啡因通过阻断腺苷受体起作用。\n\n"
        "- 要点一:半衰期五到六小时\n"
        "- 要点二:影响深度睡眠\n\n"
        "| 项目 | 数值 |\n| --- | --- |\n| 半衰期 | 5.5 小时 |\n",
        encoding="utf-8",
    )
    return p


def make_txt(tmp: Path) -> Path:
    p = tmp / "demo.txt"
    p.write_text(
        "第一段内容。\n第二行属于第一段。\n\n第二段内容,另起一段。\n\n第三段。",
        encoding="utf-8",
    )
    return p


def _tiny_png(tmp: Path, name: str = "img.png") -> Path:
    from PIL import Image

    p = tmp / name
    Image.new("RGB", (64, 48), (30, 60, 120)).save(p)
    return p


def make_docx(tmp: Path) -> Path:
    from docx import Document

    p = tmp / "demo.docx"
    doc = Document()
    doc.add_heading("咖啡因与睡眠", level=1)
    doc.add_paragraph(BODY_ZH)
    doc.add_heading("机制细节", level=2)
    doc.add_paragraph("下面是列表:", style="List Bullet")
    doc.add_paragraph("半衰期五到六小时", style="List Bullet")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "项目"
    table.rows[0].cells[1].text = "数值"
    table.rows[1].cells[0].text = "半衰期"
    table.rows[1].cells[1].text = "5.5 小时"
    doc.add_picture(str(_tiny_png(tmp)))
    doc.save(p)
    return p


def make_text_pdf(tmp: Path, pages: int = 3) -> Path:
    """文本型 PDF:每页标题行(大字号)+ 正文 + 页眉页脚(重复文本)。"""
    import pymupdf

    p = tmp / "text.pdf"
    doc = pymupdf.open()
    for i in range(pages):
        page = doc.new_page(width=595, height=842)
        page.insert_text((72, 40), "机密·内部资料", fontsize=8, fontname="china-s")  # 页眉(跨页相同,应去重)
        page.insert_text((72, 80), f"第 {i+1} 章 标题", fontsize=16, fontname="china-s")
        for j in range(4):
            page.insert_text(
                (72, 120 + j * 22),
                f"这是第 {i+1} 页的第 {j+1} 段正文,内容用于测试层级与阅读顺序。",
                fontsize=10,
                fontname="china-s",
            )
        page.insert_text((72, 800), "第 1 页 / 共 3 页", fontsize=8, fontname="china-s")  # 页脚(跨页相同,应去重)
    doc.save(p)
    doc.close()
    return p


def make_two_column_pdf(tmp: Path) -> Path:
    import pymupdf

    p = tmp / "columns.pdf"
    doc = pymupdf.open()
    page = doc.new_page(width=595, height=842)
    for j in range(5):
        page.insert_text((72, 100 + j * 22), f"左栏内容第 {j+1} 行", fontsize=10, fontname="china-s")
        page.insert_text((330, 100 + j * 22), f"右栏内容第 {j+1} 行", fontsize=10, fontname="china-s")
    doc.save(p)
    doc.close()
    return p


def make_scanned_pdf(tmp: Path) -> Path:
    """扫描型 PDF:页面只有图片,无文本层。"""
    import pymupdf
    from PIL import Image, ImageDraw

    png = tmp / "scan.png"
    img = Image.new("RGB", (1200, 800), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 100), "扫描件测试文字", fill="black")
    img.save(png)

    p = tmp / "scanned.pdf"
    doc = pymupdf.open()
    page = doc.new_page(width=595, height=842)
    page.insert_image(pymupdf.Rect(0, 0, 595, 400), filename=str(png))
    doc.save(p)
    doc.close()
    return p


def make_encrypted_pdf(tmp: Path) -> Path:
    import pymupdf

    p = tmp / "encrypted.pdf"
    doc = pymupdf.open()
    doc.new_page(width=595, height=842).insert_text((72, 100), "加密文档", fontsize=12)
    doc.save(p, encryption=pymupdf.PDF_ENCRYPT_AES_256, user_pw="secret", owner_pw="secret2")
    doc.close()
    return p


def make_malformed(tmp: Path, name: str = "bad.pdf") -> Path:
    p = tmp / name
    p.write_bytes(b"\x89PNG\r\n\x1a\n garbage not a pdf not a zip " + b"\x00" * 64)
    return p


def make_zipbomb_docx(tmp: Path, uncompressed_mb: int = 60, entries: int = 50) -> Path:
    """构造 DOCX 壳 + 巨大可压缩条目(测试 P0 限额)。"""
    p = tmp / "bomb.docx"
    with zipfile.ZipFile(p, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", "<Types/>")
        z.writestr("word/document.xml", "<w:document/>")
        z.writestr("media/big.bin", b"\x00" * (uncompressed_mb * 1024 * 1024))
        for i in range(entries):
            z.writestr(f"word/extra{i}.xml", "<x/>")
    return p


def ocr_engine_available() -> tuple[bool, str]:
    try:
        from rapidocr_onnxruntime import RapidOCR  # noqa: F401

        return True, ""
    except Exception as exc:  # noqa: BLE001
        return False, str(exc)
